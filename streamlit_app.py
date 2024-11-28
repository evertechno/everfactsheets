import streamlit as st
import google.generativeai as genai
from fpdf import FPDF
from docx import Document
from io import BytesIO

# Configure the API key securely from Streamlit's secrets
genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

# Streamlit App UI
st.title("AI Product/Service Factsheet Generator")
st.write("Use AI to generate a professional factsheet for your product or service. Fill out the fields below to customize your factsheet.")

# Input fields for the product/service
product_name = st.text_input("Product/Service Name")
description = st.text_area("Product/Service Description", height=150)
target_audience = st.text_input("Target Audience")
key_features = st.text_area("Key Features (separate by commas)", height=100)
price = st.text_input("Pricing Information")
customer_benefits = st.text_area("Customer Benefits", height=100)
unique_selling_point = st.text_input("Unique Selling Point (USP)")
competitors = st.text_area("Competitors (separate by commas)")
industry = st.text_input("Industry")
launch_date = st.date_input("Launch Date")
contact_info = st.text_area("Contact Information", height=80)
website = st.text_input("Website URL")

# Additional fields for customization
tone_of_voice = st.selectbox("Tone of Voice", ["Formal", "Casual", "Professional", "Friendly", "Technical"])
factsheet_style = st.selectbox("Factsheet Layout", ["Minimalist", "Corporate", "Creative", "Innovative"])
include_logo = st.checkbox("Include Logo in the Factsheet")
include_image = st.checkbox("Include an Image")
image_url = st.text_input("Image URL (if applicable)")

# Button to generate factsheet
if st.button("Generate Factsheet"):
    if product_name and description and target_audience and key_features:
        try:
            # Prepare the input prompt for the generative AI model
            prompt = f"Generate a product/service factsheet for a product named '{product_name}'. The description is '{description}'. " \
                     f"The target audience is '{target_audience}'. The key features are {key_features}. Pricing information: {price}. " \
                     f"Customer benefits: {customer_benefits}. Unique selling point: {unique_selling_point}. Competitors: {competitors}. " \
                     f"Industry: {industry}. Launch date: {launch_date}. Contact information: {contact_info}. Website: {website}. " \
                     f"Generate with a tone of '{tone_of_voice}' and a layout style of '{factsheet_style}'."

            # Load and configure the model
            model = genai.GenerativeModel('gemini-1.5-flash')
            
            # Generate response from the model
            response = model.generate_content(prompt)
            factsheet_text = response.text
            
            # Display generated factsheet
            st.write("Generated Factsheet:")
            st.write(factsheet_text)

            # Export options
            export_options = st.radio("Export as:", ["PDF", "DOCX", "Text"])

            # PDF Export
            if export_options == "PDF":
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                if include_logo:
                    pdf.image('logo_path.png', 10, 8, 33)  # Add logo (provide logo path)
                if include_image and image_url:
                    pdf.image(image_url, x=50, w=100)  # Add an image (provide valid URL)
                pdf.multi_cell(0, 10, factsheet_text)
                pdf_output = "/mnt/data/factsheet.pdf"
                pdf.output(pdf_output)
                st.download_button("Download PDF", pdf_output)

            # DOCX Export
            elif export_options == "DOCX":
                doc = Document()
                doc.add_heading(f"{product_name} Factsheet", 0)
                doc.add_paragraph(factsheet_text)
                if include_logo:
                    doc.add_picture('logo_path.png', width=Inches(1.0))  # Add logo
                if include_image and image_url:
                    doc.add_picture(image_url, width=Inches(2.0))  # Add image
                doc_output = "/mnt/data/factsheet.docx"
                doc.save(doc_output)
                st.download_button("Download DOCX", doc_output)

            # Text Export
            elif export_options == "Text":
                text_output = "/mnt/data/factsheet.txt"
                with open(text_output, "w") as f:
                    f.write(factsheet_text)
                st.download_button("Download Text File", text_output)

        except Exception as e:
            st.error(f"Error: {e}")
    else:
        st.error("Please fill out all the fields before generating the factsheet.")

# Feature to allow users to compare different versions of factsheets
if st.checkbox("Generate Multiple Versions for Comparison"):
    num_versions = st.slider("Number of Versions", min_value=1, max_value=5, value=3)
    
    for i in range(num_versions):
        version_prompt = f"Generate version {i+1} of the factsheet for the product '{product_name}' with the following details: {description}. " \
                         f"Include tone of '{tone_of_voice}', layout '{factsheet_style}', and focus on these features: {key_features}. " \
                         f"Make the factsheet different from the others."

        model = genai.GenerativeModel('gemini-1.5-flash')
        version_response = model.generate_content(version_prompt)
        st.write(f"Version {i+1} Generated Factsheet:")
        st.write(version_response.text)

        export_version_option = st.radio(f"Export Version {i+1} as:", ["PDF", "DOCX", "Text"])
        if export_version_option == "PDF":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, version_response.text)
            pdf_output = f"/mnt/data/factsheet_version_{i+1}.pdf"
            pdf.output(pdf_output)
            st.download_button(f"Download Version {i+1} as PDF", pdf_output)
        elif export_version_option == "DOCX":
            doc = Document()
            doc.add_heading(f"{product_name} Factsheet - Version {i+1}", 0)
            doc.add_paragraph(version_response.text)
            doc_output = f"/mnt/data/factsheet_version_{i+1}.docx"
            doc.save(doc_output)
            st.download_button(f"Download Version {i+1} as DOCX", doc_output)
        elif export_version_option == "Text":
            text_output = f"/mnt/data/factsheet_version_{i+1}.txt"
            with open(text_output, "w") as f:
                f.write(version_response.text)
            st.download_button(f"Download Version {i+1} as Text File", text_output)
